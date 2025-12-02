from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO
import re

class ReportGenrator:
    def __init__(self, theme = "", search_result = [0, 0], query = "", students_data = [], source = ["", ""], matrix_json = []):        
        # 初始化所有需要的欄位
        self.report_title = "智慧財產權實戰策略期末報告"
        self.theme = theme
        self.search_result = search_result
        self.query = query
        self.scope = self.extract_ipc_scope(query)
        self.students_data = students_data
        self.source = source
        self.matrix_json = matrix_json

    def set(self, theme = None, search_result = None, query = None, students_data = None, source = None, matrix_json = None):
        if theme:
            self.theme = theme
        if search_result:
            self.search_result = search_result
        if query:
            self.query = query
            self.scope = self.extract_ipc_scope(query)
        if students_data:
            self.students_data = students_data
        if source:
            self.source = source
        if matrix_json:
            self.matrix_json = matrix_json

    def extract_ipc_scope(self, query):
        """
        從檢索字串中擷取 IC (IPC/CPC) 範圍，自動去除括號。
        """
        # Regex 解析：
        # 1. \(?       -> 匹配開頭可選的左括號 '('
        # 2. (IC=[^)]+) -> 【捕獲群組】抓取 IC= 開頭，且內容不包含 ')' 的所有字元
        # 3. \)?       -> 匹配結尾可選的右括號 ')'
        pattern = r"\(?(IC=[^)]+)\)?"
        
        # 搜尋所有符合的片段 (通常 IPC 設定會在最後面，我們取最後一個匹配的或是特定的)
        match = re.search(pattern, query)
        
        if match:
            # group(1) 會自動排除掉括號，只回傳中間的內容
            return match.group(1).strip()
        return None


    def gen_report(self):

        def set_chinese_font(run, size=12, bold=False):
            """設定中文字型的 Helper Function"""
            run.font.name = 'Times New Roman' # 西文部分
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體') # 中文部分 (可改微軟正黑體)
            run.font.size = Pt(size)
            run.bold = bold
        
        doc = Document()
        title = doc.add_heading(level=0)
        run = title.add_run('智慧財產權實戰策略期末報告')
        set_chinese_font(run, size=20, bold=True)

        run_dot = title.add_run('.')

        run_dot.font.color.rgb = RGBColor(255, 255, 255)
        run_dot.font.size = Pt(20)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # --- 2. 成員名單表格 ---
        # 成員表格
        doc.add_heading('成員名單', level=1)
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        
        # 表頭
        headers = ['姓名', '學號']
        for i in range(2):
            text = headers[i]
            cell = table.cell(0, i)
            run = cell.paragraphs[0].add_run(text)
            set_chinese_font(run, bold=True)

            text = self.students_data[i]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell = table.cell(1, i)
            run = cell.paragraphs[0].add_run(text)
            set_chinese_font(run)

        doc.add_paragraph()

        # --- 3. 檢索流程 ---
        # 
        h1 = doc.add_heading(level=1)
        set_chinese_font(h1.add_run('檢索流程'), size=16, bold=True)

        # 選定主題
        p = doc.add_paragraph()
        set_chinese_font(p.add_run('一、選定主題：'), bold=True)
        set_chinese_font(p.add_run(f"本研究針對{self.theme}專利分析與布局進行專利檢索分析。"))

        p = doc.add_paragraph()
        set_chinese_font(p.add_run('二、專家建議、網路資源、相關報導、領導廠商（或有利人士）'), bold=True)
        # 檢索條件
        p = doc.add_paragraph()
        set_chinese_font(p.add_run('三、訂出主題關鍵字：'), bold=True)
        set_chinese_font(p.add_run('設定以下檢索條件。'))
        query_para = doc.add_paragraph(self.query)
        query_para.runs[0].font.size = Pt(9)
        query_para.runs[0].font.color.rgb = RGBColor(50, 50, 150)

        p = doc.add_paragraph()
        set_chinese_font(p.add_run("四、專利資料庫："), bold=True)
        set_chinese_font(p.add_run("全球專利檢索系統GPSS。"))
        
        p = doc.add_paragraph()
        set_chinese_font(p.add_run("五、主題相關專利："), bold=True)
        set_chinese_font(p.add_run(f"本次專利檢索共獲得 {self.search_result[0]} 筆 專利資料，經過檢索結果去重與專利家族去重處理後，最終共有 {self.search_result[1]} 筆專利資料納入分析。"))
        p = doc.add_paragraph()
        if self.source[0]:
            source_text = f"IPC或CPC國際分類號採用：參考{self.source[0]}，"
        else:
            source_text = ""
        if self.scope:
            scope_text = f"將專利檢索式之國際專利分類號設定在{self.scope}"
        else:
            scope_text = f"將專利檢索式之國際專利分類號設定在 IPC = IPC"
        set_chinese_font(p.add_run("六、國際專利分類號設定："), bold=True)
        set_chinese_font(p.add_run(source_text + scope_text))

        p = doc.add_paragraph()
        set_chinese_font(p.add_run("七、抽樣檢索："), bold=True)
        set_chinese_font(p.add_run("除了關鍵詞與國際分類號，檢索條件仍採取and or not等控制條件。"))

        charts = [
            ("技術領域分類分析"),
            ("領導者分析（誰是這個領域的領導者）"),
            ("布局戰場分析（那些國家是布局戰場）"),
            ("專利申請趨勢")
        ]
        for i, title_text in enumerate(charts):
            if i % 2 == 0:
                doc.add_page_break()
            doc.add_heading(title_text, level=2)
            doc.add_picture(f".data/chart{i + 1}.png", width=Inches(5))
            # 如果您有圖片檔案，可以使用: doc.add_picture('path_to_image.png', width=Inches(6))

        doc.add_page_break()

        if not self.matrix_json:
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer

        # --- 5. 技術與功效分類表 ---
        # 技術分類表格
        doc.add_heading('專利技術圖分析 - 技術分類', level=2)
        
        tech_data = self.matrix_json["technologies"]
        
        table = doc.add_table(rows=len(tech_data) + 1, cols=2)
        table.style = 'Table Grid'
        cell = table.cell(0, 0)
        run = cell.paragraphs[0].add_run("技術分類")
        set_chinese_font(run, size=10, bold=True)
        cell = table.cell(0, 1)
        run = cell.paragraphs[0].add_run("技術分類檢索詞")
        set_chinese_font(run, size=10, bold=True)    

        for i, tech in enumerate(tech_data):
            cell = table.cell(i + 1, 0)
            run = cell.paragraphs[0].add_run(tech["label"])
            set_chinese_font(run, size=10, bold=False)
            cell = table.cell(i + 1, 1)
            run = cell.paragraphs[0].add_run(tech["boolean"])
            set_chinese_font(run, size=10, bold=False)

        if self.source[1]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_chinese_font(p.add_run(f"資料來源：{self.source[1]}"))

        # 功效分類表格
        doc.add_heading('功效分類', level=2)
        
        eff_data = self.matrix_json["efficacies"]
        
        table = doc.add_table(rows=len(eff_data) + 1, cols=2)
        table.style = 'Table Grid'
        cell = table.cell(0, 0)
        run = cell.paragraphs[0].add_run("功效分類")
        set_chinese_font(run, size=10, bold=True)
        cell = table.cell(0, 1)
        run = cell.paragraphs[0].add_run("功效分類檢索詞")
        set_chinese_font(run, size=10, bold=True) 
        for i, eff in enumerate(eff_data):
            cell = table.cell(i + 1, 0)
            run = cell.paragraphs[0].add_run(eff["label"])
            set_chinese_font(run, size=10, bold=False)
            cell = table.cell(i + 1, 1)
            run = cell.paragraphs[0].add_run(eff["boolean"])
            set_chinese_font(run, size=10, bold=False)

        if self.source[1]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_chinese_font(p.add_run(f"資料來源：{self.source[1]}"))
        # --- 6. 技術功效矩陣分析 ---
        # 
        doc.add_heading('三、技術功效矩陣分析', level=1)
        doc.add_picture(f".data/chart5.png", width=Inches(5))
        
        # 儲存
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer